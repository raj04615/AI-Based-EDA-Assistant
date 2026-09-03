"""
Pytest Suite for Multi-Format Extraction, Magic Byte Signature Detection, and ZIP Batch Failure Isolation.
"""

import os
import io
import tempfile
import zipfile
import pytest
from fastapi.testclient import TestClient

import database
import ingestion
import config
from app import app

client = TestClient(app)


def test_detect_file_type_magic_bytes():
    """Verify format detection via magic byte signatures."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(b"%PDF-1.4 sample pdf content")
        pdf_path = f.name
        
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
        f.write(b"Just plain text data")
        txt_path = f.name

    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
        f.write(b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR...")
        png_path = f.name

    try:
        assert ingestion.detect_file_type(pdf_path) == "pdf"
        assert ingestion.detect_file_type(txt_path) == "txt"
        assert ingestion.detect_file_type(png_path) == "image"
    finally:
        os.remove(pdf_path)
        os.remove(txt_path)
        os.remove(png_path)


def test_extract_txt_md():
    """Test text and Markdown extractor."""
    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", encoding="utf-8", delete=False) as f:
        f.write("# Quarterly Analysis\n\nRevenue grew by 20% in Q3 due to cloud product sales.")
        md_path = f.name

    try:
        pages_data = ingestion.extract_txt_md(md_path)
        assert len(pages_data) == 1
        assert "Quarterly Analysis" in pages_data[0]["text"]
        assert pages_data[0]["page_label"] == "Page 1"
    finally:
        os.remove(md_path)


def test_extract_docx():
    """Test Microsoft Word (.docx) extractor."""
    import docx
    doc = docx.Document()
    doc.add_heading("Project Report", level=1)
    doc.add_paragraph("This is an introductory paragraph describing financial gains.")
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Profit Margin"
    table.cell(1, 1).text = "24%"

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        docx_path = f.name

    doc.save(docx_path)

    try:
        sections = ingestion.extract_docx(docx_path)
        assert len(sections) >= 1
        assert "Project Report" in sections[0]["text"]
        assert len(sections[0]["tables"]) == 1
        assert "Profit Margin" in sections[0]["tables"][0]
    finally:
        os.remove(docx_path)


def test_extract_csv_excel():
    """Test CSV & Excel data extractors into Markdown tables."""
    import pandas as pd
    df = pd.DataFrame({"Department": ["Sales", "Engineering"], "Budget": ["$5M", "$12M"]})
    
    with tempfile.NamedTemporaryFile(suffix=".csv", mode="w", encoding="utf-8", delete=False) as f:
        df.to_csv(f.name, index=False)
        csv_path = f.name

    try:
        data = ingestion.extract_excel_csv(csv_path, "csv")
        assert len(data) == 1
        assert len(data[0]["tables"]) == 1
        assert "Engineering" in data[0]["tables"][0]
        assert "Sheet 'Data'" in data[0]["page_label"]
    finally:
        os.remove(csv_path)


def test_non_utf8_csv_extraction():
    """Verify CSV with non-UTF8 characters (e.g. 0xa0 non-breaking space / latin-1) ingests cleanly."""
    raw_bytes = b"ID,Name,Value\n1,Item \xa0 A,100\n2,Caf\xe9,200\n"
    with tempfile.NamedTemporaryFile(suffix=".csv", mode="wb", delete=False) as f:
        f.write(raw_bytes)
        csv_path = f.name

    try:
        data = ingestion.extract_excel_csv(csv_path, "csv")
        assert len(data) == 1
        assert len(data[0]["tables"]) == 1
        assert "Item" in data[0]["tables"][0]
    finally:
        os.remove(csv_path)


def test_zip_upload_batch_and_failure_isolation():
    """Test ZIP upload containing mixed valid files, unsupported format, and nested ZIP."""
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, "w") as z:
        # 1. Valid TXT file
        z.writestr("notes.txt", "Analytical findings for Q4 2026.")
        
        # 2. Unsupported file (.exe)
        z.writestr("binary.exe", b"\x4d\x5a\x90\x00\x03\x00")
        
        # 3. Nested ZIP archive
        nested_buf = io.BytesIO()
        with zipfile.ZipFile(nested_buf, "w") as nz:
            nz.writestr("inner.txt", "hello")
        z.writestr("nested.zip", nested_buf.getvalue())

    zip_bytes = zip_buffer.getvalue()
    files = {"file": ("test_archive.zip", zip_bytes, "application/zip")}

    response = client.post("/upload", files=files)
    assert response.status_code == 200
    
    res_data = response.json()
    assert res_data["batch"] is True
    assert len(res_data["documents"]) == 1  # notes.txt accepted
    assert len(res_data["skipped"]) == 2    # binary.exe & nested.zip skipped

    valid_doc = res_data["documents"][0]
    assert valid_doc["filename"] == "notes.txt"
    assert valid_doc["file_type"] == "txt"
    assert valid_doc["status"] == "processing"

    # Clean up DB entry
    database.delete_document(valid_doc["doc_id"])
